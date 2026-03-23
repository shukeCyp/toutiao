"""
ArticleDownloader - 文章下载客户端
使用 Playwright 打开文章页面，提取正文内容，生成 docx 文档（含图片）
"""

import os
import re
import io
import asyncio
import random
import requests
from html.parser import HTMLParser
from datetime import datetime
from PIL import Image
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from logger import get_logger
from fingerprint import random_fingerprint

log = get_logger('downloader')


def safe_filename(name, max_len=80):
    """将标题转为安全的文件名"""
    # 去除不安全字符
    name = re.sub(r'[\\/:*?"<>|\n\r\t]', '', name)
    name = name.strip('. ')
    if len(name) > max_len:
        name = name[:max_len]
    return name or 'untitled'


# 署名行匹配模式：如 "文| 杨磊"、"编辑丨姜召"、"责编：张三"、"来源：xxx" 等
_SIGNATURE_PATTERN = re.compile(
    r'^[\s]*(文|作者|编辑|责编|责任编辑|审核|初审|终审|复审|校对|排版|图片|来源|供稿|记者|通讯员|本文作者|撰文|策划|监制|出品)'
    r'[\s]*[|丨/／:：]'
)


class ArticleContentParser(HTMLParser):
    """解析 article-content HTML，提取正文段落和图片（跳过标题和元信息）"""

    def __init__(self):
        super().__init__()
        self.elements = []  # [{'type': 'text', 'text': ...}, {'type': 'image', 'url': ...}]
        self._current_tag = None
        self._in_article = False
        self._text_buf = ''
        self._skip_tags = {'script', 'style', 'noscript'}
        self._skip_depth = 0
        self._skip_block_depth = 0  # 跳过 h1 / article-meta 等块级元素

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag in self._skip_tags:
            self._skip_depth += 1
            return

        # 跳过 <h1> 标题
        if tag == 'h1':
            self._flush_text()
            self._skip_block_depth += 1
            return

        # 跳过 <div class="article-meta"> 元信息
        if tag == 'div' and 'article-meta' in attrs_dict.get('class', ''):
            self._flush_text()
            self._skip_block_depth += 1
            return

        # 在跳过块内，追踪嵌套深度
        if self._skip_block_depth > 0:
            if tag in ('div', 'h1'):
                self._skip_block_depth += 1
            return

        if tag == 'p':
            self._flush_text()
            self._current_tag = 'p'

        elif tag == 'img':
            self._flush_text()
            # 优先 data-src，fallback src
            url = attrs_dict.get('data-src', '') or attrs_dict.get('src', '')
            # 过滤 base64 占位图
            if url and not url.startswith('data:'):
                self.elements.append({'type': 'image', 'url': url})

    def handle_endtag(self, tag):
        if tag in self._skip_tags and self._skip_depth > 0:
            self._skip_depth -= 1
            return

        # 处理跳过块的关闭
        if self._skip_block_depth > 0:
            if tag in ('h1', 'div'):
                self._skip_block_depth -= 1
            return

        if tag == 'p':
            self._flush_text()
            self._current_tag = None

    def handle_data(self, data):
        if self._skip_depth > 0 or self._skip_block_depth > 0:
            return
        text = data.strip()
        if text:
            self._text_buf += text

    def _flush_text(self):
        text = self._text_buf.strip()
        self._text_buf = ''
        if not text:
            return
        # 清理 XML 不兼容的控制字符
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        if not text:
            return
        # 过滤署名行
        if _SIGNATURE_PATTERN.match(text):
            return
        self.elements.append({'type': 'text', 'text': text})


class HuiwenContentParser(HTMLParser):
    """
    解析 feed.huiwen.co 文章 HTML（.yl-content 结构）
    正文在 <section> > <span> 中，图片在 <div class="yl-img-wrapper"> > <img> 中，
    底部 <p> 标签中的署名行需过滤。
    """

    def __init__(self):
        super().__init__()
        self.elements = []
        self._text_buf = ''
        self._skip_tags = {'script', 'style', 'noscript'}
        self._skip_depth = 0
        self._in_section = False
        self._section_depth = 0

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag in self._skip_tags:
            self._skip_depth += 1
            return

        if self._skip_depth > 0:
            return

        if tag == 'section':
            if not self._in_section:
                self._flush_text()
                self._in_section = True
                self._section_depth = 1
            else:
                self._section_depth += 1

        elif tag == 'p':
            self._flush_text()

        elif tag == 'img':
            self._flush_text()
            url = attrs_dict.get('data-src', '') or attrs_dict.get('src', '')
            if url and not url.startswith('data:'):
                if url.startswith('//'):
                    url = 'https:' + url
                self.elements.append({'type': 'image', 'url': url})

    def handle_endtag(self, tag):
        if tag in self._skip_tags and self._skip_depth > 0:
            self._skip_depth -= 1
            return

        if self._skip_depth > 0:
            return

        if tag == 'section' and self._in_section:
            self._section_depth -= 1
            if self._section_depth <= 0:
                self._flush_text()
                self._in_section = False
                self._section_depth = 0

        elif tag == 'p':
            self._flush_text()

    def handle_data(self, data):
        if self._skip_depth > 0:
            return
        text = data.strip()
        if text:
            self._text_buf += text

    def _flush_text(self):
        text = self._text_buf.strip()
        self._text_buf = ''
        if not text:
            return
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        if not text:
            return
        if _SIGNATURE_PATTERN.match(text):
            return
        # 过滤 "声明：" 提示文字
        if text.startswith('声明：') or text.startswith('声明:'):
            return
        self.elements.append({'type': 'text', 'text': text})


def parse_article_html(html):
    """解析头条文章 HTML 内容"""
    parser = ArticleContentParser()
    parser.feed(html)
    parser._flush_text()
    return parser.elements


def parse_huiwen_html(html):
    """解析 feed.huiwen.co 文章 HTML 内容"""
    parser = HuiwenContentParser()
    parser.feed(html)
    parser._flush_text()
    return parser.elements


def _is_huiwen_url(url):
    """判断是否为 feed.huiwen.co 域名的文章"""
    return 'huiwen.co' in url or 'feed.huiwen' in url


def _is_people_url(url):
    """判断是否为 m2.people.cn 域名的文章"""
    return 'm2.people.cn' in url or 'm.people.cn' in url


def _get_source_type(url):
    """根据 URL 判断文章来源类型，返回 'huiwen' / 'people' / 'toutiao'"""
    if _is_huiwen_url(url):
        return 'huiwen'
    if _is_people_url(url):
        return 'people'
    return 'toutiao'


_SUPPORTED_DOMAINS = [
    'toutiao.com',
    'huiwen.co',
    'm2.people.cn',
    'm.people.cn',
]


def is_supported_url(url):
    """判断 URL 是否为已适配的域名，不支持的域名无法正确抓取"""
    if not url:
        return False
    return any(d in url for d in _SUPPORTED_DOMAINS)


def _get_referer(source_type):
    """根据来源类型返回对应的 Referer"""
    return {
        'huiwen': 'https://feed.huiwen.co/',
        'people': 'http://m2.people.cn/',
        'toutiao': 'https://www.toutiao.com/',
    }.get(source_type, 'https://www.toutiao.com/')


def _parse_proxy_line(line):
    """
    解析单行代理字符串，返回 Playwright proxy dict
    格式: IP:端口 用户名 密码
    """
    parts = line.strip().split()
    if not parts:
        return None
    server = parts[0]
    if not server.startswith('http'):
        server = 'http://' + server
    proxy = {'server': server}
    if len(parts) >= 3:
        proxy['username'] = parts[1]
        proxy['password'] = parts[2]
    return proxy


def _check_proxy(proxy_config, timeout=8):
    """通过代理发送请求检测是否可用，返回 True/False"""
    try:
        server = proxy_config['server']
        proxies = {'http': server, 'https': server}
        if proxy_config.get('username'):
            auth_server = server.replace('://', f'://{proxy_config["username"]}:{proxy_config["password"]}@')
            proxies = {'http': auth_server, 'https': auth_server}
        resp = requests.get('http://httpbin.org/ip', proxies=proxies, timeout=timeout)
        return resp.status_code == 200
    except Exception:
        return False


def _random_proxy(proxy_pool_text):
    """从多行代理文本中随机选取一个可用的代理，返回 Playwright proxy dict 或 None"""
    if not proxy_pool_text or not proxy_pool_text.strip():
        return None
    lines = [l.strip() for l in proxy_pool_text.strip().splitlines() if l.strip()]
    if not lines:
        return None
    random.shuffle(lines)
    for line in lines:
        proxy = _parse_proxy_line(line)
        if not proxy:
            continue
        log.info(f'检测代理: {proxy["server"]}')
        if _check_proxy(proxy):
            log.info(f'代理可用: {proxy["server"]}')
            return proxy
        log.warning(f'代理不可用: {proxy["server"]}，尝试下一个')
    log.warning('所有代理均不可用，将不使用代理')
    return None


def download_image_bytes(url, referer='https://www.toutiao.com/', timeout=15):
    """下载图片并返回 bytes"""
    try:
        if url.startswith('//'):
            url = 'https:' + url
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
            'Referer': referer,
        }
        resp = requests.get(url, headers=headers, timeout=timeout)
        if resp.status_code == 200 and len(resp.content) > 500:
            return resp.content
    except Exception as e:
        log.debug(f'下载图片失败: {url[:80]} -> {e}')
    return None


def crop_watermark(img_bytes):
    """
    裁剪图片右下角水印
    从底部裁掉约 8% 的高度，去除头条水印区域
    """
    try:
        img = Image.open(io.BytesIO(img_bytes))
        w, h = img.size

        # 图片太小则不裁剪
        if h < 100 or w < 100:
            return img_bytes

        # 裁掉底部 8% 的高度
        crop_h = int(h * 0.08)
        cropped = img.crop((0, 0, w, h - crop_h))

        buf = io.BytesIO()
        # 保持原格式，默认 PNG
        fmt = img.format or 'PNG'
        cropped.save(buf, format=fmt)
        buf.seek(0)
        return buf.getvalue()
    except Exception as e:
        log.debug(f'裁剪水印失败: {e}')
        return img_bytes


def read_docx_elements(doc_path):
    """
    从本地 docx 文件读取内容，返回 elements 列表（文字 + 图片）

    Args:
        doc_path: docx 文件路径

    Returns:
        list: [{'type': 'text', 'text': ...}, {'type': 'image', 'data': bytes}, ...]
    """
    log.info(f'读取本地 docx: {doc_path}')
    doc = Document(doc_path)
    elements = []

    for para in doc.paragraphs:
        # 检查段落中是否包含图片
        has_image = False
        for run in para.runs:
            # 查找 inline shape（图片）
            drawing_elements = run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
            ) or run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'
            )
            if not drawing_elements:
                # 也检查 drawingML
                drawing_elements = run._element.findall(
                    './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
                )

            for drawing in drawing_elements:
                # 提取 blip（图片引用）
                blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId:
                        try:
                            image_part = doc.part.related_parts[rId]
                            elements.append({'type': 'image', 'data': image_part.blob})
                            has_image = True
                        except (KeyError, Exception) as e:
                            log.debug(f'读取图片失败: rId={rId}, err={e}')

        # 如果段落不含图片，且有文字内容，作为文本元素
        if not has_image:
            text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', para.text.strip())
            if text and not _SIGNATURE_PATTERN.match(text):
                elements.append({'type': 'text', 'text': text})

    log.info(f'从 docx 读取到 {len(elements)} 个元素')
    return elements


def _clean_text(text):
    """清理文本中 XML 不兼容的控制字符（NULL、退格等）"""
    # 保留 tab(\x09)、换行(\x0A)、回车(\x0D)，去掉其余控制字符
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)


def generate_docx(elements, save_path, source_url=''):
    """
    根据解析后的元素列表生成纯文本 docx 文件（宋体）

    Args:
        elements: [{'type': 'text'|'image', ...}]
            image 元素支持两种：
            - {'type': 'image', 'url': '...'} — 从网络下载
            - {'type': 'image', 'data': bytes} — 直接使用本地数据
        save_path: 保存文件完整路径
        source_url: 原文链接，用于确定图片下载 Referer 和是否裁水印
    """
    doc = Document()

    # 设置默认字体为宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    source_type = _get_source_type(source_url) if source_url else 'toutiao'
    referer = _get_referer(source_type)
    need_crop = (source_type == 'toutiao')

    for elem in elements:
        if elem['type'] == 'text':
            doc.add_paragraph(_clean_text(elem['text']))

        elif elem['type'] == 'image':
            img_bytes = elem.get('data')  # 本地图片数据
            if not img_bytes and elem.get('url'):
                img_bytes = download_image_bytes(elem['url'], referer=referer)
                if img_bytes and need_crop:
                    img_bytes = crop_watermark(img_bytes)

            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                except Exception as e:
                    log.debug(f'插入图片失败: {e}')
                    doc.add_paragraph('[图片加载失败]')
            else:
                doc.add_paragraph('[图片加载失败]')

    # 确保目录存在
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    log.info(f'文档已保存: {save_path}')
    return save_path


async def fetch_article_elements(article_url, headless=True, proxy_pool=''):
    """
    打开文章页面，提取正文元素列表（文字段落 + 图片）
    自动识别 toutiao.com 和 feed.huiwen.co 使用不同的选择器和解析器

    Args:
        article_url: 文章 URL
        headless: 是否使用无头模式
        proxy_pool: 代理池文本（每行一个代理），为空则不使用代理

    Returns:
        list: [{'type': 'text', 'text': ...}, {'type': 'image', 'url': ...}, ...]
    """
    source_type = _get_source_type(article_url)
    selector_map = {
        'huiwen': '.yl-content',
        'people': 'article',
        'toutiao': '.article-content',
    }
    content_selector = selector_map[source_type]
    log.info(f'提取文章内容: {article_url} (headless={headless}, type={source_type})')

    proxy_config = _random_proxy(proxy_pool)
    fp = random_fingerprint()
    pw = await async_playwright().start()
    launch_kwargs = {
        'headless': headless,
        'args': [
            '--disable-blink-features=AutomationControlled',
            '--no-first-run',
            '--disable-infobars',
        ],
    }
    if proxy_config:
        launch_kwargs['proxy'] = proxy_config
    browser = await pw.chromium.launch(**launch_kwargs)

    try:
        context = await browser.new_context(
            viewport=fp['viewport'],
            user_agent=fp['user_agent'],
            locale=fp['locale'],
            timezone_id=fp['timezone_id'],
            color_scheme=fp['color_scheme'],
            device_scale_factor=fp['device_scale_factor'],
            extra_http_headers=fp['extra_http_headers'],
        )
        page = await context.new_page()

        await page.goto(article_url, wait_until='load', timeout=30000)

        # people.cn 通过 AJAX 动态加载内容，需等待 JS 执行完成
        if source_type == 'people':
            await page.wait_for_timeout(3000)

        try:
            await page.wait_for_selector(content_selector, timeout=10000)
        except Exception:
            log.warning(f'未找到 {content_selector}，尝试等待更久')
            await page.wait_for_timeout(3000)

        # 滚动到底部触发懒加载图片
        await page.evaluate('window.scrollTo(0, document.body.scrollHeight)')
        await page.wait_for_timeout(2000)

        selector_js = content_selector.replace("'", "\\'")
        html = await page.evaluate(f'''
            () => {{
                const el = document.querySelector('{selector_js}');
                return el ? el.innerHTML : '';
            }}
        ''')

        await context.close()

    finally:
        await browser.close()
        await pw.stop()

    if not html:
        raise RuntimeError('未能提取到文章内容')

    log.info(f'HTML 内容长度: {len(html)}')
    parser_map = {
        'huiwen': parse_huiwen_html,
        'people': parse_article_html,
        'toutiao': parse_article_html,
    }
    elements = parser_map[source_type](html)
    log.info(f'解析到 {len(elements)} 个元素')
    return elements


class ArticleDownloader:
    """文章下载器：提取内容 -> 生成 docx"""

    async def download(self, article_url, save_dir, category='', title='', headless=True, proxy_pool=''):
        """
        下载单篇文章为 docx

        Args:
            article_url: 文章 URL
            save_dir: 根保存目录
            category: 文章分类（作为子文件夹）
            title: 文章标题（作为文件名，为空则从页面提取）
            headless: 是否使用无头模式
            proxy_pool: 代理池文本

        Returns:
            str: 保存的文件路径
        """
        elements = await fetch_article_elements(article_url, headless=headless, proxy_pool=proxy_pool)

        # 如果没有传入标题，使用默认名称
        if not title:
            title = 'untitled_' + datetime.now().strftime('%Y%m%d%H%M%S')

        # 构建保存路径: save_dir/category/title.docx
        filename = safe_filename(title) + '.docx'
        if category:
            folder = os.path.join(save_dir, safe_filename(category))
        else:
            folder = save_dir
        save_path = os.path.join(folder, filename)

        # 生成 docx
        generate_docx(elements, save_path, source_url=article_url)

        return save_path
